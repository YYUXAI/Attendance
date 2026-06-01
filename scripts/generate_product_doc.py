# -*- coding: utf-8 -*-
"""生成考勤机器人产品说明 Word 文档（修复中文乱码）。"""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
# 英文文件名，避免部分环境打开路径乱码
OUT_EN = DOCS / "Attendance_Bot_Product_Guide.docx"
OUT_CN = DOCS / "考勤机器人产品功能说明.docx"

# 宋体在 Windows Word 上兼容性最好
FONT_ASCII = "Times New Roman"
FONT_CN = "宋体"


def _set_run_font(run, *, size: Pt | None = None, bold: bool = False) -> None:
    run.font.name = FONT_ASCII
    run.bold = bold
    if size is not None:
        run.font.size = size
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), FONT_ASCII)
    rFonts.set(qn("w:hAnsi"), FONT_ASCII)
    rFonts.set(qn("w:eastAsia"), FONT_CN)
    rFonts.set(qn("w:cs"), FONT_CN)


def _set_paragraph_font(paragraph) -> None:
    for run in paragraph.runs:
        _set_run_font(run)


def _add_run(paragraph, text: str, *, size: Pt | None = Pt(11), bold: bool = False):
    run = paragraph.add_run(text)
    _set_run_font(run, size=size, bold=bold)
    return run


def _para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    _add_run(p, text)


def _bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    _add_run(p, text)


def _heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(level=level)
    p.clear()
    sizes = {1: Pt(16), 2: Pt(14), 3: Pt(12)}
    _add_run(p, text, size=sizes.get(level, Pt(12)), bold=True)


def _cell_text(cell, text: str) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    _add_run(p, text, size=Pt(10))


def _patch_document_styles(doc: Document) -> None:
    for name in ("Normal", "List Bullet", "List Number"):
        try:
            style = doc.styles[name]
        except KeyError:
            continue
        style.font.name = FONT_ASCII
        style.font.size = Pt(11)
        rPr = style.element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn("w:ascii"), FONT_ASCII)
        rFonts.set(qn("w:hAnsi"), FONT_ASCII)
        rFonts.set(qn("w:eastAsia"), FONT_CN)
        rFonts.set(qn("w:cs"), FONT_CN)
    for lvl in range(1, 4):
        try:
            style = doc.styles[f"Heading {lvl}"]
        except KeyError:
            continue
        style.font.name = FONT_ASCII
        rPr = style.element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn("w:eastAsia"), FONT_CN)


def build() -> tuple[Path, Path]:
    DOCS.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)
    _patch_document_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(title, "考勤 Telegram 机器人 — 产品功能说明", size=Pt(18), bold=True)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(
        sub,
        f"文档版本：V1.0    更新日期：{date.today().isoformat()}",
        size=Pt(10),
    )
    doc.add_paragraph()

    _para(
        doc,
        "本文档面向产品、运营与管理人员，用通俗语言说明当前系统已具备的主要能力、"
        "使用场景与规则边界，不涉及技术实现细节。",
    )

    _heading(doc, "一、截图打卡（AI 识别）", 1)
    _para(
        doc,
        "使用场景：员工在本人班次的考勤群中，发送一张包含 TIME.IS 网页大钟、"
        "以及 Slack 个人资料浮窗的截图，由机器人自动识别并完成打卡记录。",
    )

    _heading(doc, "1.1 员工怎么操作", 2)
    _bullet(doc, "在考勤群发送截图（支持照片或图片文件）。")
    _bullet(doc, "机器人提示：正在识别打卡截图，请稍候（约 1 分钟）。")
    _bullet(doc, "识别通过后，机器人回复打卡成功信息；未通过则提示失败原因。")

    _heading(doc, "1.2 系统校验什么", 2)
    _bullet(
        doc,
        "是不是本人：从截图中的 Slack 浮窗读取用户名，必须与员工在系统中登记的 "
        "Telegram 用户名一致；不一致则拒绝打卡，用于防止代打卡。",
    )
    _bullet(
        doc,
        "时间对不对：从整张截图中收集所有类似时钟的时间（含主钟、页脚城市时间等）；"
        "若出现多个时间，选取与员工发图时刻最接近、且相差不超过 1 小时的一个。",
    )
    _bullet(
        doc,
        "时区规则：时间比对统一按北京时间，不按班次配置中的其他时区（如曼谷）换算。",
    )
    _bullet(
        doc,
        "读不出主钟：不使用消息发送时间代替截图时间；识别失败则打卡失败。",
    )

    _heading(doc, "1.3 用户看到的反馈", 2)
    _bullet(doc, "失败时仅两类提示：打卡失败：姓名不一致。或 打卡失败：时间不一致。")
    _bullet(
        doc,
        "成功时展示：英文名、工号、部门、班次、时区、打卡时间、时间来源（截图 AI 识别）、"
        "截图用户校验结果、图片编号等。",
    )

    _heading(doc, "1.4 说明", 2)
    _para(
        doc,
        "当前识别方式为纯文字识别（OCR），不依赖大模型看图理解。"
        "截图需尽量清晰，主钟与 Slack 浮窗完整可见，且建议使用当前时刻的实时截图。",
    )

    _heading(doc, "二、报备休息（请假）", 1)
    _para(doc, "使用场景：员工计划休假，在私聊机器人中提交申请，由上级审批。")

    _heading(doc, "2.1 操作流程", 2)
    _bullet(doc, "私聊机器人，点击主菜单「报备休息」。")
    _bullet(doc, "选择休假类型：年假、病假、事假等，也支持自定义类型。")
    _bullet(doc, "填写休假起止日期。")
    _bullet(doc, "确认信息后提交。")

    _heading(doc, "2.2 提交之后", 2)
    _bullet(doc, "申请进入审批流程，审批人会在私聊中收到待办通知。")
    _bullet(doc, "审批通过后，系统记录有效休假日期，后续考勤统计会纳入休假状态。")

    _heading(doc, "三、离岗报备（短时间离开工位）", 1)
    _para(
        doc,
        "使用场景：员工在工作时段内需要短暂离开工位（如外出办事），"
        "提前报备并获得审批，避免被记为异常缺勤。",
    )

    _heading(doc, "3.1 操作流程", 2)
    _bullet(doc, "私聊机器人，点击「离岗报备」。")
    _bullet(doc, "填写预计离开时间与预计返回时间。")
    _bullet(doc, "确认后提交。")

    _heading(doc, "3.2 提交之后", 2)
    _bullet(doc, "同样进入审批流程，审批人私聊处理通过或驳回。")
    _bullet(
        doc,
        "审批通过后，系统会在对应时间点自动处理生效、结束、过期，无需人工逐条改状态。",
    )

    _heading(doc, "四、审批人功能", 1)
    _para(doc, "使用场景：班组长、负责人等在私聊中处理下属提交的休假或离岗申请。")

    _heading(doc, "4.1 可处理事项", 2)
    _bullet(doc, "休假申请：通过或驳回，可填写审批备注。")
    _bullet(doc, "离岗申请：通过或驳回。")

    _heading(doc, "4.2 处理结果", 2)
    _bullet(doc, "申请人会收到审批结果通知（通过或驳回及说明）。")
    _bullet(doc, "休假若审批通过，可向相关考勤群发送公告（视配置与业务规则）。")

    _heading(doc, "五、质检（QC）流程", 1)
    _para(
        doc,
        "使用场景：按班次计划对员工进行抽查，要求在规定时间内完成指定手势的锁屏拍照并提交。",
    )

    _heading(doc, "5.1 流程概览", 2)
    _bullet(doc, "自动开轮：到设定时间，系统在考勤群发布「本轮质检开始」公告。")
    _bullet(
        doc,
        "私信通知：被抽中的员工收到私信，说明需在锁屏界面按示例手势拍照，并提示确认或取消。",
    )
    _bullet(doc, "员工操作：确认开始，上传照片，二次确认提交；也可取消后重新上传。")
    _bullet(doc, "超时规则：规定时间内未完成，记为本轮质检失败。")
    _bullet(doc, "收尾汇总：轮次或班次结束后，在群里发布汇总公告（展示每人各轮结果）。")
    _bullet(doc, "群聊中提供部分收尾类按钮操作，供管理员或流程配合使用。")

    _heading(doc, "5.2 特别说明", 2)
    _para(
        doc,
        "部分历史公告文案中提到可向机器人发送 /qc 指令重启流程；"
        "当前产品版本中尚未提供该命令入口，实际推进依赖系统自动私信通知与轮询任务。",
    )

    _heading(doc, "六、系统自动执行（无需员工点击）", 1)

    _heading(doc, "6.1 考勤相关提醒", 2)
    _bullet(
        doc,
        "班次开始后：向考勤群发送班次开始类通知；班组长、部门负责人可能收到私信，"
        "内容包括应到岗人数、已到岗人数、未到岗人员名单等。",
    )
    _bullet(doc, "班次结束后：向考勤群发送下班打卡提醒。")
    _bullet(
        doc,
        "统计逻辑会结合打卡记录、已批准休假、已生效离岗等信息，减少误报。",
    )

    _heading(doc, "6.2 其它后台任务", 2)
    _bullet(doc, "审批单自动派发给对应审批人。")
    _bullet(doc, "各类 Telegram 消息通过队列异步发送，避免阻塞主流程。")
    _bullet(doc, "离岗记录按时间自动生效、结束或过期。")
    _bullet(doc, "质检：自动开轮、私信通知、超时判定、轮次收尾与班次汇总等。")

    _heading(doc, "七、功能一览（便于评审）", 1)

    table = doc.add_table(rows=8, cols=3)
    table.style = "Table Grid"
    headers = ["模块", "谁在用", "一句话说明"]
    rows = [
        ("截图打卡", "全体员工", "群里发 TIME.IS+Slack 图，校验本人且时间在 1 小时内（北京时间）"),
        ("报备休息", "全体员工", "私聊申请休假，审批通过后计入休假"),
        ("离岗报备", "全体员工", "私聊申请短时离岗，到点自动生效/结束"),
        ("审批", "班组长/负责人", "私聊通过或驳回休假、离岗"),
        ("质检 QC", "被抽查员工", "私信拍照确认，超时记失败，最后群发汇总"),
        ("自动提醒", "全员/管理者", "上下班点到点发群消息与私信统计"),
        ("注册与信息", "新员工", "私聊注册、查看本人考勤概况"),
    ]
    for i, h in enumerate(headers):
        _cell_text(table.rows[0].cells[i], h)
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            _cell_text(table.rows[r_idx].cells[c_idx], val)

    doc.add_paragraph()
    _para(
        doc,
        "文档结束。如有规则调整（如打卡时间容差、时区策略），以系统配置与后续版本说明为准。",
    )

    # 全文再扫一遍，防止遗漏未设置东亚字体的 run
    for p in doc.paragraphs:
        _set_paragraph_font(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _set_paragraph_font(p)

    doc.save(OUT_EN)
    doc.save(OUT_CN)
    return OUT_EN, OUT_CN


if __name__ == "__main__":
    en, cn = build()
    # 校验：读回第一段中文是否正常
    check = Document(en)
    sample = check.paragraphs[0].text
    print("saved:", en)
    print("saved:", cn)
    print("sample:", sample[:40])
